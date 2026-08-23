"""PR-CONTEXT -- build PORTABILITYRISK_MASTER_REPORT_V1.docx from the verified Markdown.

The DOCX is generated FROM the final Markdown so the two cannot carry conflicting numbers.
Markdown is parsed rather than retyped.

Layout handling:
  * US Letter portrait; wide tables (>6 columns) are placed in their own LANDSCAPE section
    and the document returns to portrait afterwards.
  * A refreshable TOC field, so headings drive the contents page.
  * Numbered headings via the built-in Heading styles.
  * Header with the project name, footer with a PAGE field.
  * `code spans` become monospaced runs; *italics* stay italic, which carries the scientific
    names; **bold** stays bold.
  * Table and figure captions with sequential numbering.
  * Tables get `cantSplit` on every row so no row is torn across a page boundary, and
    headings get keepNext so none is stranded at a page bottom.
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1A, 0x1A, 0x1A)
ACC = RGBColor(0x0B, 0x6F, 0xA4)
MONO = "Consolas"


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def add_field(par, instr):
    r = par.add_run()
    r._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    t = OxmlElement("w:instrText")
    t.set(qn("xml:space"), "preserve")
    t.text = instr
    r._r.append(t)
    r._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    r2 = par.add_run("field")
    r2._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))
    return par


def keep_next(par):
    p = par._p.get_or_add_pPr()
    p.append(_el("w:keepNext", **{"w:val": "1"}))
    p.append(_el("w:keepLines", **{"w:val": "1"}))


def cant_split(row):
    tr = row._tr.get_or_add_trPr()
    tr.append(_el("w:cantSplit"))


def repeat_header(row):
    tr = row._tr.get_or_add_trPr()
    tr.append(_el("w:tblHeader", **{"w:val": "true"}))


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_runs(par, text, base_size=9.5):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = MONO; r.font.size = Pt(base_size - 1.0)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1]); r.italic = True
        else:
            r = par.add_run(tok)
        r.font.size = Pt(base_size)


def set_orientation(section, landscape):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = Inches(11), Inches(8.5)
        section.left_margin = section.right_margin = Inches(0.7)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = Inches(8.5), Inches(11)
        section.left_margin = section.right_margin = Inches(1.0)
    section.top_margin = section.bottom_margin = Inches(0.9)


def build(md_path, out_path, figdir):
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    st.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    set_orientation(sec, False)
    hdr = sec.header.paragraphs[0]
    hdr.text = "PortabilityRisk — Master Technical and Scientific Report  |  V1.0"
    hdr.runs[0].font.size = Pt(8)
    hdr.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ftr = sec.footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = ftr.add_run("Page ")
    rr.font.size = Pt(8)
    add_field(ftr, "PAGE")
    rr2 = ftr.add_run(" of ")
    rr2.font.size = Pt(8)
    add_field(ftr, "NUMPAGES")
    for r in ftr.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    # ---------------- title page
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PortabilityRisk"); r.bold = True; r.font.size = Pt(30)
    r.font.color.rgb = ACC
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Master Technical and Scientific Report")
    r.font.size = Pt(17); r.font.color.rgb = INK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Direct replicon evidence for the portability of acquired antimicrobial "
                  "resistance genes in ESKAPE Gram-negative pathogens")
    r.italic = True; r.font.size = Pt(11.5)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version 1.0   ·   21 August 2026"); r.font.size = Pt(11); r.bold = True
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Analysis status: PORTABILITY_CONTEXT_COMPLETE")
    r.font.size = Pt(10); r.font.name = MONO; r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x52)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Internal scientific record  ·  authorship placeholder")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------------- TOC
    h = doc.add_heading("Table of contents", level=1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    p = doc.add_paragraph()
    r = p.add_run("This contents page is a live field. In Word, select all and press F9 "
                  "(or right-click → Update Field) to refresh page numbers.")
    r.italic = True; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    tbl_n, fig_n = 0, 0
    last_heading, last_label = "", ""
    i = 0
    landscape_now = False
    n = len(lines)
    while i < n:
        ln = lines[i]
        s = ln.strip()

        if not s or s == "---":
            i += 1
            continue

        # ---- image
        m = re.match(r"^!\[(.*?)\]\((.+?)\)$", s)
        if m:
            path = os.path.join(figdir, os.path.basename(m.group(2)))
            if os.path.exists(path):
                fig_n += 1
                pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                width = Inches(9.2 if landscape_now else 6.3)
                pp.add_run().add_picture(path, width=width)
                cap = lines[i + 2].strip() if i + 2 < n else ""
                cm = re.match(r"^\*\*Figure \d+\.\*\*\s*(.*)$", cap)
                capt = doc.add_paragraph()
                capt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = capt.add_run("Figure %d. " % fig_n); rr.bold = True
                rr.font.size = Pt(9)
                add_runs(capt, cm.group(1) if cm else "", 9)
                if cm:
                    i += 3
                    continue
            i += 1
            continue

        # ---- heading
        if s.startswith("#"):
            lvl = len(s) - len(s.lstrip("#"))
            txt = s.lstrip("# ").strip()
            if lvl == 1 and txt.startswith("PortabilityRisk —"):
                i += 1
                continue
            last_heading = re.sub(r"[*`]", "", txt)
            last_label = ""
            hh = doc.add_heading(level=min(lvl, 3))
            for tok in INLINE.split(txt):
                if not tok:
                    continue
                if tok.startswith("*") and tok.endswith("*") and not tok.startswith("**"):
                    rr = hh.add_run(tok.strip("*")); rr.italic = True
                else:
                    hh.add_run(tok.replace("**", ""))
            keep_next(hh)
            i += 1
            continue

        # ---- table
        if s.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            rows = [[c.strip() for c in r.strip("|").split("|")] for r in block
                    if not re.match(r"^\|[\s:|-]+\|$", r)]
            if not rows:
                continue
            ncol = max(len(r) for r in rows)
            rows = [r + [""] * (ncol - len(r)) for r in rows]
            wide = ncol > 6
            if wide != landscape_now:
                ns = doc.add_section(WD_SECTION.NEW_PAGE)
                set_orientation(ns, wide)
                landscape_now = wide
            tbl_n += 1
            title = last_label or last_heading
            title = re.sub(r"^[A-Z](\.\d+)?\.?\s+", "", title).strip()
            cap = doc.add_paragraph()
            rr = cap.add_run("Table %d. " % tbl_n); rr.bold = True; rr.font.size = Pt(9)
            rr.font.color.rgb = ACC
            rt = cap.add_run(title if title else "Summary")
            rt.font.size = Pt(9); rt.italic = True
            rt.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            keep_next(cap)
            last_label = ""
            t = doc.add_table(rows=len(rows), cols=ncol)
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = True
            for ri, row in enumerate(rows):
                cant_split(t.rows[ri])
                if ri == 0:
                    repeat_header(t.rows[ri])
                for ci, cell in enumerate(row):
                    c = t.cell(ri, ci)
                    c.text = ""
                    par = c.paragraphs[0]
                    par.paragraph_format.space_after = Pt(1)
                    par.paragraph_format.space_before = Pt(1)
                    add_runs(par, cell, 8.0 if ncol > 5 else 8.6)
                    if ri == 0:
                        for rn in par.runs:
                            rn.bold = True
            doc.add_paragraph()
            continue

        # ---- blockquote block
        if s.startswith(">"):
            qb = []
            while i < n and lines[i].strip().startswith(">"):
                qb.append(lines[i].strip().lstrip(">").strip())
                i += 1
            for q in qb:
                if not q:
                    continue
                if q.startswith("#"):
                    pq = doc.add_paragraph()
                    rq = pq.add_run(q.lstrip("# ").strip())
                    rq.bold = True; rq.font.size = Pt(11); rq.font.color.rgb = ACC
                    keep_next(pq)
                    continue
                pq = doc.add_paragraph()
                pq.paragraph_format.left_indent = Inches(0.3)
                pq.paragraph_format.space_after = Pt(4)
                add_runs(pq, re.sub(r"^\d+\.\s*", "", q), 9.5)
                for rn in pq.runs:
                    rn.font.color.rgb = RGBColor(0x22, 0x44, 0x55)
            continue

        # ---- code fence
        if s.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            for c in code:
                pc = doc.add_paragraph()
                pc.paragraph_format.left_indent = Inches(0.25)
                pc.paragraph_format.space_after = Pt(0)
                rc = pc.add_run(c if c.strip() else " ")
                rc.font.name = MONO; rc.font.size = Pt(8.2)
                rc.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            doc.add_paragraph()
            continue

        # ---- list
        if re.match(r"^[-*]\s+", s) or re.match(r"^\d+\.\s+", s):
            style = "List Bullet" if re.match(r"^[-*]\s+", s) else "List Number"
            while i < n and (re.match(r"^[-*]\s+", lines[i].strip())
                             or re.match(r"^\d+\.\s+", lines[i].strip())):
                txt = re.sub(r"^([-*]|\d+\.)\s+", "", lines[i].strip())
                pl = doc.add_paragraph(style=style)
                add_runs(pl, txt, 9.8)
                i += 1
            continue

        # ---- paragraph
        para = [s]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(#|\||>|```|!\[|[-*]\s|\d+\.\s|---)", lines[i].strip()):
            para.append(lines[i].strip())
            i += 1
        joined = " ".join(para)
        if (joined.startswith("**") and joined.endswith("**")
                and joined.count("**") == 2 and len(joined) < 90):
            last_label = joined.strip("*")
        pp = doc.add_paragraph()
        add_runs(pp, joined, 10)

    if landscape_now:
        ns = doc.add_section(WD_SECTION.NEW_PAGE)
        set_orientation(ns, False)
    doc.save(out_path)
    return tbl_n, fig_n


if __name__ == "__main__":
    t, fg = build(sys.argv[1], sys.argv[2], sys.argv[3])
    import hashlib
    p = sys.argv[2]
    print("wrote %s" % p)
    print("  bytes  %d" % os.path.getsize(p))
    print("  tables %d" % t)
    print("  figures %d" % fg)
    print("  sha256 %s" % hashlib.sha256(open(p, "rb").read()).hexdigest())
